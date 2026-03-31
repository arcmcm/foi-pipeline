#!/usr/bin/env python3
"""
foi_research_prototype.py

Prototype pipeline:
- Reads local documents in ./downloads/
- Supports:
  - Local .doc / .docx (Harvey-style table)  -> extracts claims, and links to a converted PDF in ./downloads_pdf/
  - Local .pdf (extracts text + number claims)
- Outputs CSV files:
  - claims.csv (truncated context)
  - claims_full.csv (full context)

Notes:
- PDF metric inference is intentionally looser than DOCX so narrative reports can produce claims.
- We apply plausibility filters to reduce noise from years/page numbers/etc.
"""

from __future__ import annotations

import csv
import io
import json
import os
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

import pdfplumber
from docx import Document


# -----------------------------
# Data models
# -----------------------------

@dataclass
class DocumentRecord:
    source: str          # file path on disk (e.g. downloads/Foo.docx)
    kind: str            # "doc" | "docx" | "pdf" | "unknown"
    text: str


@dataclass
class Claim:
    source: str
    metric: str
    value: int
    context: str
    context_page: int | None = None


# -----------------------------
# DOC/DOCX parsing (Harvey-style)
# -----------------------------

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # tables (important for Harvey FOI docs)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_t = (cell.text or "").strip()
                cell_t = re.sub(r"\s+", " ", cell_t)
                if cell_t:
                    row_text.append(cell_t)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


# -----------------------------
# PDF parsing
# -----------------------------

def extract_text_from_pdf(path: str) -> str:
    with open(path, "rb") as f:
        data = f.read()
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        return "\n".join((page.extract_text() or "") for page in pdf.pages)


# -----------------------------
# Claim extraction helpers
# -----------------------------

def normalize_space(s: str) -> str:
    s = re.sub(r"\s+", " ", (s or "")).strip()

    # Fix DOCX/PDF extraction artifact: "at17" -> "at 17"
    s = re.sub(r"\bat(?=\d)", "at ", s)

    return s


def is_plausible_homeless_count_docx(n: int) -> bool:
    if n < 0:
        return False
    if n in (0, 1, 2, 3):
        return True
    if n < 10_000_000:
        return True
    return False


def is_plausible_homeless_count_pdf(n: int) -> bool:
    if n < 0:
        return False

    # exclude likely years
    if 1900 <= n <= 2100:
        return False

    # exclude obvious tiny noise that appears everywhere
    if n in (4, 5, 6, 7, 8, 9, 10):
        return False

    if n == 0:
        return True
    if 1 <= n <= 5_000_000:
        return True

    return False


def infer_metric_from_text(text: str) -> Optional[str]:
    low = text.lower()

    # Order matters: more specific before less specific (Harvey-style)
    if "single person applications" in low and "family applications" in low:
        return "homeless_applications_by_household_type"
    if "four hurdles" in low:
        return "homeless_applications_passed_four_hurdles"
    if "registered as homeless" in low:
        return "households_registered_homeless"
    if "temporary accommodation units" in low:
        return "temporary_accommodation_units"
    if 'how many "homeless applications"' in low or "how many homeless applications" in low:
        return "homeless_applications"

    # PDF / narrative fallback metrics (less strict)
    if "temporary accommodation" in low:
        return "temporary_accommodation_general"
    if "homelessness" in low or "homeless" in low:
        return "homelessness_general"
    if "asylum" in low:
        return "asylum_related_housing_general"
    if "hotel" in low:
        return "hotel_accommodation_general"
    if "housing emergency" in low or "housing" in low:
        return "housing_emergency_general"

    return None


def extract_numbers(text: str) -> List[int]:
    """
    Extract integers with basic cleanup:
    - allows comma separators: 12,345
    - ignores decimals
    """
    nums: List[int] = []
    for m in re.finditer(r"\b\d{1,3}(?:,\d{3})*\b", text):
        raw = m.group(0)
        n = int(raw.replace(",", ""))
        nums.append(n)
    return nums


def extract_claims_from_docx(source: str, text: str) -> List[Claim]:
    metric_doc = infer_metric_from_text(text) or "unknown_docx"
    claims: List[Claim] = []

    lines = [normalize_space(x) for x in text.splitlines()]
    lines = [x for x in lines if x]

    topical_re = re.compile(r"\b(homeless|homelessness|temporary accommodation|bed and breakfast|hotel|housing)\b", re.I)

    for line in lines:
        if not topical_re.search(line):
            continue

        # DOCX tables: ignore question-only lines. Prefer table rows that look like answers.
        if "|" not in line:
            continue
        answerish = (":" in line) or (" as at" in line.lower()) or ("number of" in line.lower())
        if not answerish:
            continue

        # Skip rates/shares
        if "%" in line:
            continue

        # Skip financial years like 2022/23, 2024/25
        if re.search(r"\b20\d{2}/\d{2}\b", line):
            continue

        metric = infer_metric_from_text(line) or metric_doc

        for n in extract_numbers(line):
            if not is_plausible_homeless_count_docx(n):
                continue
            ctx = line[:240]
            claims.append(Claim(source=source, metric=metric, value=n, context=ctx))

    return claims


def extract_claims_from_pdf(source: str, text: str) -> List[Claim]:
    """
    Page-aware PDF strategy:
    - Open PDF with pdfplumber
    - For each page:
      - Split into lines
      - topical gate
      - skip % and financial-year patterns and ratio-ish lines
      - infer metric
      - extract numbers with plausibility filters
      - attach context_page (1-based)
    """
    claims: List[Claim] = []

    topical_re = re.compile(r"\b(homeless|homelessness|temporary accommodation|asylum|hotel|housing emergency|housing)\b", re.I)

    with pdfplumber.open(source) as pdf:
        for page_index, page in enumerate(pdf.pages):
            page_no = page_index + 1
            page_text = page.extract_text() or ""
            lines = [normalize_space(x) for x in page_text.splitlines()]
            lines = [x for x in lines if x]

            for line in lines:
                if not topical_re.search(line):
                    continue

                if "%" in line:
                    continue

                if re.search(r"\b20\d{2}/\d{2}\b", line):
                    continue

                if "for every" in line.lower():
                    continue

                metric = infer_metric_from_text(line) or infer_metric_from_text(page_text) or "unknown_pdf"

                for n in extract_numbers(line):
                    if not is_plausible_homeless_count_pdf(n):
                        continue
                    ctx = line[:240]
                    claims.append(Claim(source=source, metric=metric, value=n, context=ctx, context_page=page_no))

    return claims


# -----------------------------
# DOC/DOCX -> PDF conversion
# -----------------------------

def ensure_word_pdfs(download_dir: str = "downloads", out_dir: str = "downloads_pdf") -> None:
    """
    Convert all .doc and .docx in download_dir to PDF into out_dir.
    Overwrites PDFs if they already exist.
    """
    src = Path(download_dir)
    out = Path(out_dir)
    if not src.is_dir():
        return
    out.mkdir(parents=True, exist_ok=True)

    docs = list(src.glob("*.doc")) + list(src.glob("*.docx"))
    if not docs:
        return

    for p in docs:
        # LibreOffice command
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out),
                str(p),
            ],
            check=False,
        )


def _word_source_to_pdf_relpath(source_path: str) -> str:
    """
    If source_path points at downloads/Foo.doc or downloads/Foo.docx,
    return downloads_pdf/Foo.pdf (relative path used in Sheets links).
    Otherwise return original.
    """
    if not source_path:
        return source_path
    p = Path(source_path)
    if p.suffix.lower() not in [".doc", ".docx", ".docx"]:
        return source_path
    return str(Path("downloads_pdf") / (p.stem + ".pdf"))


# -----------------------------
# CSV output
# -----------------------------

def write_claims_csv(path: str, claims: List[Claim], max_context_chars: int | None = 200) -> None:
    # UTF-8 with BOM helps Excel/Sheets open cleanly
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["document_title", "source", "source_page_url", "metric", "value", "context_page", "context"])

        for c in claims:
            # Make links use relative paths (downloads/... or downloads_pdf/...) so your Sheets hyperlink can prefix http://100.109.150.79:8000/
            source_rel = c.source

            # If claim is from Word, point to the converted PDF instead
            if str(source_rel).lower().endswith((".doc", ".docx")):
                source_rel = _word_source_to_pdf_relpath(source_rel)

            document_title = os.path.basename(source_rel)

            context = c.context or ""
            context = str(context)
            if max_context_chars is not None and len(context) > max_context_chars:
                context = context[:max_context_chars].rstrip() + "…"

            source_page_url = ""
            # Only PDFs can have #page= links in our current pipeline
            if str(source_rel).lower().endswith(".pdf"):
                if c.context_page is not None:
                    source_page_url = f"{source_rel}#page={c.context_page}"
                else:
                    # for Word->PDF conversions we don't know a page number, so link to the file
                    source_page_url = source_rel

            w.writerow([document_title, source_rel, source_page_url, c.metric, c.value, c.context_page, context])


# -----------------------------
# Pipeline
# -----------------------------

def load_documents_from_downloads(download_dir: str = "downloads") -> List[DocumentRecord]:
    docs: List[DocumentRecord] = []
    if not os.path.isdir(download_dir):
        return docs

    for name in sorted(os.listdir(download_dir)):
        path = os.path.join(download_dir, name)
        low = name.lower()

        if low.endswith(".docx"):
            text = extract_text_from_docx(path)
            docs.append(DocumentRecord(source=path, kind="docx", text=text))
        elif low.endswith(".doc"):
            # Treat .doc same as .docx for linking purposes, but python-docx can't read .doc.
            # We'll rely on the converted PDF for opening; extraction won't run for .doc.
            # (If you want extraction from .doc, convert it to .docx first.)
            continue
        elif low.endswith(".pdf"):
            text = extract_text_from_pdf(path)
            docs.append(DocumentRecord(source=path, kind="pdf", text=text))
        else:
            continue

    return docs


def run() -> Dict[str, int]:
    # 1) Ensure PDFs exist for Word files (so Sheets links can open them)
    ensure_word_pdfs("downloads", "downloads_pdf")

    # 2) Load docs and extract claims
    docs = load_documents_from_downloads("downloads")

    claims: List[Claim] = []
    for d in docs:
        if d.kind == "docx":
            claims.extend(extract_claims_from_docx(d.source, d.text))
        elif d.kind == "pdf":
            claims.extend(extract_claims_from_pdf(d.source, d.text))

    findings: List[dict] = []

    write_claims_csv("claims.csv", claims)
    write_claims_csv("claims_full.csv", claims, max_context_chars=None)

    return {
        "document_count": len(docs),
        "claim_count": len(claims),
        "finding_count": len(findings),
    }


if __name__ == "__main__":
    print(json.dumps(run(), indent=2))
